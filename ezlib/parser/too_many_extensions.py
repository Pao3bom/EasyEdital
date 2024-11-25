import os
import shutil
import hashlib
import subprocess
from striprtf.striprtf import rtf_to_text
from docx import Document as DocxDocument
from fitz import Document as PdfDocument # PyMuPDF for PDF handling
from bs4 import BeautifulSoup
import aiofiles



async def parse_txt(file_path):
    # Encodings we can try
    encodings = ['utf-8', 'utf-16', 'cp1252', 'windows-1252', 'latin1',]

    for encoding in encodings:
        try:
            async with aiofiles.open(file_path, 'r', encoding=encoding) as file:
                return await file.read()
        except UnicodeDecodeError:
            # If decoding fails, try the next encoding
            continue
    
    # If all encodings fail, raise an exception
    raise ValueError(f"Unable to decode file at {file_path} with any of the given encodings.")



async def parse_html(file_path):
    # Encodings we can try
    encodings = ['utf-8', 'utf-16', 'cp1252', 'windows-1252', 'latin1',]
    
    for encoding in encodings:
        try:
            async with aiofiles.open(file_path, 'r', encoding=encoding) as file:
                html_content = await file.read()
                soup = BeautifulSoup(html_content, 'html.parser')
                return soup.get_text()
        except UnicodeDecodeError as e:
            print(f"[WARNING] Failed to decode {file_path} with encoding {encoding}: {e}")
            continue
    
    # If all encodings fail, raise an error
    raise ValueError(f"Unable to decode the file {file_path} with tried encodings: {encodings}")




def parse_pdf(file_path):
    """Parse text from a PDF file."""
    text = ""
    with PdfDocument(file_path) as doc:
        for page in doc:
            text += page.get_text("text")
    return text




def parse_docx(file_path):
    return "\n".join(p.text for p in DocxDocument(file_path).paragraphs)





def parse_rtf(file_path):
    """Parse RTF files using striprtf with encoding detection."""
    try:
        # Attempt to read the file with utf-8 encoding
        with open(file_path, 'r', encoding='utf-8') as file:
            rtf_content = file.read()
    except UnicodeDecodeError:
        # Fallback to latin1 or another encoding
        with open(file_path, 'r', encoding='latin1') as file:
            rtf_content = file.read()
    
    return rtf_to_text(rtf_content)




def parse_doc(filepath, retry_limit=3, temp_dir=None):
    # Convert PosixPath to string if necessary
    filepath = str(filepath)

    # Ensure the file exists
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")
    
    # Hash the file path to avoid conflicts
    hash = hashlib.md5(filepath.encode()).hexdigest()

    # Create a subdirectory for temporary files
    if temp_dir is None:
        temp_dir = os.path.join(os.path.dirname(filepath), f"_TEMP_{hash}")
    else:
        temp_dir = os.path.join(temp_dir, f"_TEMP_{hash}")
    os.makedirs(temp_dir, exist_ok=True)

    temp_doc_path = os.path.join(temp_dir, os.path.basename(filepath))
    predicted_docx_path = os.path.splitext(temp_doc_path)[0] + ".docx"

    attempts = 0

    while attempts <= retry_limit:
        try:
            # Copy the original file to the temporary directory
            shutil.copy(filepath, temp_doc_path)

            # Convert .doc to .docx using LibreOffice
            result = subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'docx', temp_doc_path, '--outdir', temp_dir],
                check=False,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )

            # Check if conversion was successful
            if result.returncode != 0 or not os.path.exists(predicted_docx_path):
                raise RuntimeError(
                    f"LibreOffice failed for file `{filepath}`.\n"
                    f"Command: {' '.join(result.args)}\n"
                    f"Error: {result.stderr.decode()}"
                )

            # Extract text from the converted .docx file
            doc = DocxDocument(predicted_docx_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            return text

        except Exception as e:
            attempts += 1
            if attempts > retry_limit:
                raise RuntimeError(f"Failed to extract text from `{filepath}` after {retry_limit + 1} attempts.") from e
        finally:
            # Cleanup temporary files in the directory for each attempt
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir, ignore_errors=True)
