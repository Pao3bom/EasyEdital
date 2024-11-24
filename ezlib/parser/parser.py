import os
import asyncio
import aiofiles
from fitz import Document  # PyMuPDF for PDF handling
import pypandoc
from docx import Document as DocxDocument
import pytesseract
from pdf2image import convert_from_path
from bs4 import BeautifulSoup
import subprocess
from striprtf.striprtf import rtf_to_text
from tempfile import NamedTemporaryFile



async def parse_text(file_path):
    _, ext = os.path.splitext(file_path)
    ext = ext.lower().strip()

    match ext:
        case '.txt':
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                return await file.read()

        case '.pdf':
            return await asyncio.to_thread(parse_pdf, file_path)

        case '.docx':
            return await asyncio.to_thread(
                lambda: "\n".join(p.text for p in DocxDocument(file_path).paragraphs)
            )

        case '.doc':
            return extract_text_from_doc(file_path)

        case '.rtf':
            return await asyncio.to_thread(parse_rtf_with_striprtf, file_path)

        case '.html':
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                html_content = await file.read()
                soup = BeautifulSoup(html_content, 'html.parser')
                return soup.get_text()

        case '.json' | '.csv' | '.png' | '.jpg' | '.htm' | '.download' | "" | ".css":
            return ""

        case _:
            raise ValueError(f"Unsupported file type: {ext}")



def parse_rtf_with_striprtf(file_path):
    """Parse RTF files using striprtf with encoding detection."""
    try:
        # Attempt to read the file with utf-8 encoding
        with open(file_path, 'r', encoding='utf-8') as file:
            rtf_content = file.read()
    except UnicodeDecodeError:
        # Fallback to latin1 or another encoding
        with open(file_path, 'r', encoding='latin1') as file:
            rtf_content = file.read()
    
    return rtf_to_text(rtf_content)


import shutil
import hashlib




def extract_text_from_doc(filepath, retry_limit=5):
    # Convert PosixPath to string if necessary
    filepath = str(filepath)

    # Ensure the file exists
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")
    
    # Hash the file path to avoid conflicts
    hash = hashlib.md5(filepath.encode()).hexdigest()
    temp_dir = os.path.join(os.path.dirname(filepath), f"_TEMP_{hash}")
    os.makedirs(temp_dir, exist_ok=True)

    temp_doc_path = os.path.join(temp_dir, os.path.basename(filepath))
    predicted_docx_path = os.path.splitext(temp_doc_path)[0] + ".docx"

    attempts = 0

    while attempts <= retry_limit:
        try:
            # Copy the original file to the temporary directory
            shutil.copy(filepath, temp_doc_path)
            print(f"Attempt {attempts + 1}: Converting {temp_doc_path}")

            # Convert .doc to .docx using LibreOffice
            result = subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'docx', temp_doc_path, '--outdir', temp_dir],
                check=False,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )

            # Check if conversion was successful
            if result.returncode != 0 or not os.path.exists(predicted_docx_path):
                raise RuntimeError(
                    f"LibreOffice failed for file `{filepath}`.\n"
                    f"Command: {' '.join(result.args)}\n"
                    f"Error: {result.stderr.decode()}"
                )

            # Extract text from the converted .docx file
            doc = DocxDocument(predicted_docx_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            print("Conversion and extraction successful!")
            return text

        except Exception as e:
            print(f"Error during attempt {attempts + 1} for `{filepath}`: {e}")
            attempts += 1
            if attempts > retry_limit:
                raise RuntimeError(f"Failed to extract text from `{filepath}` after {retry_limit + 1} attempts.") from e
        finally:
            # Cleanup temporary files in the directory for each attempt
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir, ignore_errors=True)







def parse_pdf(file_path):
    """Parse text from a PDF file."""
    text = ""
    with Document(file_path) as doc:
        for page in doc:
            text += page.get_text("text")
    return text




def is_pdf(file_path):
    _, ext = os.path.splitext(file_path)
    ext = ext.lower().strip()
    return ext == '.pdf'


def scan_page(image, page_num):
    """OCR process for a single PDF page."""
    page_text = pytesseract.image_to_string(image)
    return f"--- Page {page_num} ---\n{page_text}\n"


async def scan_pdf(file_path, executor=None):
    images = await asyncio.to_thread(lambda: convert_from_path(file_path))

    loop = asyncio.get_running_loop()
    tasks = [
        loop.run_in_executor(executor, scan_page, image, i + 1)
        for i, image in enumerate(images)
    ]
    
    results = await asyncio.gather(*tasks)
    return "".join(results)



async def hard_parse(file_path, executor=None):
    # First try to parse the file as text
    text = await parse_text(file_path)
    
    # If no text was extracted, try to scan the file
    if not text and is_pdf(file_path):
        text = await scan_pdf(file_path, executor)
        
    return text